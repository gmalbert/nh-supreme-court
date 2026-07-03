"""
Export oral argument transcripts to PDF and DOCX.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional


def export_transcript_to_pdf(
    transcript: dict,
    output_path: Path,
    include_metadata: bool = True,
) -> bool:
    """
    Export transcript to PDF with formatting.
    
    Args:
        transcript: Transcript dict with turns
        output_path: Path to save PDF
        include_metadata: Whether to include case metadata
    
    Returns:
        True if successful
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
        
        # Create PDF
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            "TranscriptTitle",
            parent=styles["Heading1"],
            fontSize=16,
            textColor="darkblue",
            spaceAfter=20,
            alignment=TA_CENTER,
        )
        
        speaker_style = ParagraphStyle(
            "Speaker",
            parent=styles["BodyText"],
            fontSize=11,
            textColor="darkblue",
            bold=True,
            spaceAfter=4,
        )
        
        text_style = ParagraphStyle(
            "TranscriptText",
            parent=styles["BodyText"],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
        )
        
        # Build story
        story = []
        
        # Title
        case_name = transcript.get("case_name", "Oral Argument Transcript")
        story.append(Paragraph(case_name, title_style))
        
        # Metadata
        if include_metadata:
            case_number = transcript.get("case_number", "")
            argument_date = transcript.get("argument_date", "")
            
            if case_number:
                story.append(Paragraph(f"Case No. {case_number}", styles["Normal"]))
            if argument_date:
                story.append(Paragraph(f"Argued: {argument_date}", styles["Normal"]))
            
            story.append(Spacer(1, 0.5 * inch))
        
        # Transcript turns
        for turn in transcript.get("turns", []):
            speaker = turn.get("speaker", "UNKNOWN")
            text = turn.get("text", "")
            
            # Speaker
            story.append(Paragraph(f"{speaker}:", speaker_style))
            
            # Text
            story.append(Paragraph(text, text_style))
        
        # Build PDF
        doc.build(story)
        
        return True
    
    except ImportError:
        print("reportlab not installed. Install with: pip install reportlab")
        return False
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return False


def export_transcript_to_docx(
    transcript: dict,
    output_path: Path,
    include_metadata: bool = True,
) -> bool:
    """
    Export transcript to DOCX with formatting.
    
    Args:
        transcript: Transcript dict with turns
        output_path: Path to save DOCX
        include_metadata: Whether to include case metadata
    
    Returns:
        True if successful
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document
        doc = Document()
        
        # Title
        case_name = transcript.get("case_name", "Oral Argument Transcript")
        title = doc.add_heading(case_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        if include_metadata:
            case_number = transcript.get("case_number", "")
            argument_date = transcript.get("argument_date", "")
            
            if case_number:
                p = doc.add_paragraph(f"Case No. {case_number}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if argument_date:
                p = doc.add_paragraph(f"Argued: {argument_date}")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Blank line
        
        # Transcript turns
        for turn in transcript.get("turns", []):
            speaker = turn.get("speaker", "UNKNOWN")
            text = turn.get("text", "")
            
            # Speaker
            p = doc.add_paragraph()
            run = p.add_run(f"{speaker}:")
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 139)
            
            # Text
            p = doc.add_paragraph(text)
            p.style = "Body Text"
        
        # Save
        doc.save(str(output_path))
        
        return True
    
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return False
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return False


def export_transcript_to_text(
    transcript: dict,
    output_path: Path,
    include_metadata: bool = True,
) -> bool:
    """
    Export transcript to plain text.
    
    Args:
        transcript: Transcript dict with turns
        output_path: Path to save text file
        include_metadata: Whether to include case metadata
    
    Returns:
        True if successful
    """
    try:
        lines = []
        
        # Title
        case_name = transcript.get("case_name", "Oral Argument Transcript")
        lines.append(case_name.upper())
        lines.append("=" * len(case_name))
        lines.append("")
        
        # Metadata
        if include_metadata:
            case_number = transcript.get("case_number", "")
            argument_date = transcript.get("argument_date", "")
            
            if case_number:
                lines.append(f"Case No. {case_number}")
            if argument_date:
                lines.append(f"Argued: {argument_date}")
            
            lines.append("")
            lines.append("-" * 60)
            lines.append("")
        
        # Transcript turns
        for turn in transcript.get("turns", []):
            speaker = turn.get("speaker", "UNKNOWN")
            text = turn.get("text", "")
            
            lines.append(f"{speaker}:")
            lines.append(text)
            lines.append("")
        
        # Write to file
        output_path.write_text("\n".join(lines), encoding="utf-8")
        
        return True
    
    except Exception as e:
        print(f"Error generating text: {e}")
        return False
